import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF提取文本"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"PDF文本提取失败 {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """从DOCX提取文本"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.info(f"从DOCX提取了 {len(text)} 字符的文本: {file_path}")
            if not text.strip():
                logger.warning(f"DOCX文档似乎是空的: {file_path}")
                
            return text
        except Exception as e:
            logger.error(f"DOCX文本提取失败 {file_path}: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """从TXT提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"TXT文本提取失败 {file_path}: {e}")
                return ""
        except Exception as e:
            logger.error(f"TXT文本提取失败 {file_path}: {e}")
            return ""
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取文档元数据"""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "filename": path.name,
            "file_path": str(path),
            "file_size": stat.st_size,
            "created_time": stat.st_ctime,
            "modified_time": stat.st_mtime,
            "file_extension": path.suffix.lower(),
            "category": self._categorize_document(path.name)
        }
    
    def _categorize_document(self, filename: str) -> str:
        """根据文件名推断文档类别"""
        filename_lower = filename.lower()
        
        if any(keyword in filename_lower for keyword in ['hr', '人力', '员工', '薪资', '招聘', '年假', '休假', '假期']):
            return "hr"
        elif any(keyword in filename_lower for keyword in ['tech', '技术', '开发', 'api', '架构']):
            return "tech"
        elif any(keyword in filename_lower for keyword in ['policy', '政策', '制度', '规定']):
            return "policy"
        elif any(keyword in filename_lower for keyword in ['faq', '常见问题', '问答']):
            return "faq"
        else:
            return "general"
    
    def process_document(self, file_path: str) -> List[Dict[str, Any]]:
        """处理单个文档，返回文档块列表"""
        try:
            logger.info(f"开始处理文档: {file_path}")
            
            # 提取元数据
            metadata = self.extract_metadata(file_path)
            logger.info(f"文档元数据: 文件名={metadata['filename']}, 类别={metadata['category']}, 大小={metadata['file_size']}字节")
            
            # 根据文件类型提取文本
            file_extension = metadata["file_extension"]
            logger.info(f"文档类型: {file_extension}")
            
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"不支持的文件类型: {file_extension}")
                return []
            
            # 检查提取的文本
            text_length = len(text)
            if not text.strip():
                logger.warning(f"文档为空: {file_path}")
                return []
            else:
                logger.info(f"成功提取文本，长度: {text_length} 字符")
                logger.info(f"文本预览: {text[:200]}...")
            
            # 文本切分
            chunks = self.text_splitter.split_text(text)
            logger.info(f"文本已切分为 {len(chunks)} 个块，平均块大小: {text_length/max(1, len(chunks)):.1f} 字符")
            
            # 构建文档块
            document_chunks = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_id": i,
                    "chunk_text": chunk,
                    "chunk_length": len(chunk)
                })
                document_chunks.append(chunk_metadata)
            
            logger.info(f"文档处理完成: {file_path}, 生成 {len(chunks)} 个文档块")
            return document_chunks
            
        except Exception as e:
            logger.error(f"文档处理失败 {file_path}: {e}")
            return []
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """处理目录下的所有文档"""
        all_chunks = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"目录不存在: {directory_path}")
            return []
        
        # 支持的文件类型
        supported_extensions = {'.pdf', '.docx', '.txt'}
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                chunks = self.process_document(str(file_path))
                all_chunks.extend(chunks)
        
        logger.info(f"目录处理完成: {directory_path}, 总共 {len(all_chunks)} 个文档块")
        return all_chunks